import os
import pandas as pd
import pytesseract
from PIL import Image
from core.converter_factory import ConverterFactory
import tempfile
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from docxcompose.composer import Composer
from xhtml2pdf import pisa
import openpyxl
from openpyxl import Workbook


def docx_to_pdf(docx_path, **kwargs):
    """Convert DOCX to PDF using external library"""
    output_path = kwargs.get('output_path')

    try:
        from docx2pdf import convert
        convert(docx_path, output_path)
        return output_path
    except ImportError:
        raise ImportError("docx2pdf library is required for DOCX to PDF conversion")


def html_to_pdf(html_input, output_path):
    """Convert HTML to PDF using xhtml2pdf"""
    try:
        # Load HTML content
        if os.path.isfile(html_input):
            with open(html_input, 'r', encoding='utf-8') as html_file:
                html_content = html_file.read()
        else:
            html_content = html_input

        # Create PDF
        with open(output_path, "w+b") as result_file:
            pisa_status = pisa.CreatePDF(
                src=html_content,
                dest=result_file,
                encoding='utf-8'
            )

        return not pisa_status.err
    except Exception as e:
        print(f"HTML to PDF conversion failed: {str(e)}")
        return False


def excel_to_pdf(excel_path, **kwargs):
    """Convert Excel to PDF using xhtml2pdf"""
    sheet_name = kwargs.get('sheet_name')
    output_path = kwargs.get('output_path')

    try:
        # Read the Excel file
        df = pd.read_excel(excel_path, sheet_name=sheet_name) if sheet_name else pd.read_excel(excel_path)

        if df.empty:
            raise ValueError("Excel sheet is empty.")

        # Convert DataFrame to HTML
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                table {{ border-collapse: collapse; width: 100%; font-family: Arial, sans-serif; }}
                th, td {{ border: 1px solid #ddd; text-align: left; padding: 8px; }}
                th {{ background-color: #f2f2f2; }}
                tr:nth-child(even) {{ background-color: #f9f9f9; }}
            </style>
        </head>
        <body>{df.to_html(index=False)}</body>
        </html>
        """

        # Convert HTML to PDF
        with open(output_path, "w+b") as result_file:
            pisa_status = pisa.CreatePDF(src=html_content, dest=result_file, encoding='utf-8')

        if pisa_status.err:
            raise ValueError(f"PDF generation failed")

        return output_path
    except Exception as e:
        raise ValueError(f"Error converting Excel to PDF: {str(e)}")


def pdf_to_docx(pdf_path, output_path=None, chunk_size=3):
    """Convert PDF to DOCX using a chunking approach"""
    try:
        # Set output path if not provided
        if output_path is None:
            output_path = str(Path(pdf_path).with_suffix('.docx'))

        # Create temp directory for chunks
        temp_dir = tempfile.mkdtemp()
        chunk_files = []

        # Read the PDF
        reader = PdfReader(pdf_path)
        total_pages = len(reader.pages)

        # Process the PDF in chunks
        for start_page in range(0, total_pages, chunk_size):
            doc = Document()
            end_page = min(start_page + chunk_size, total_pages)

            # Process pages in this chunk
            for page_num in range(start_page, end_page):
                text = reader.pages[page_num].extract_text()
                doc.add_heading(f"Page {page_num + 1}", level=2)
                doc.add_paragraph(text)

            # Save this chunk
            chunk_path = os.path.join(temp_dir, f"chunk_{start_page}.docx")
            doc.save(chunk_path)
            chunk_files.append(chunk_path)

        # Merge all chunks
        if chunk_files:
            master = Document(chunk_files[0])
            composer = Composer(master)

            for chunk_path in chunk_files[1:]:
                doc = Document(chunk_path)
                composer.append(doc)

            composer.save(output_path)
        else:
            Document().save(output_path)

        # Clean up temporary files
        for chunk_path in chunk_files:
            try:
                os.remove(chunk_path)
            except Exception:
                pass
        try:
            os.rmdir(temp_dir)
        except Exception:
            pass

        return output_path
    except Exception as e:
        raise ValueError(f"PDF to DOCX conversion failed: {str(e)}")


def create_csv_from_excel(excel_path, **kwargs):
    """Convert Excel to CSV"""
    sheet_name = kwargs.get('sheet_name')
    output_path = kwargs.get('output_path')
    all_sheets = kwargs.get('all_sheets', False)

    try:
        if all_sheets:
            # Convert all sheets
            excel = pd.ExcelFile(excel_path)
            base_name = os.path.splitext(output_path)[0]
            output_files = []

            for sheet in excel.sheet_names:
                df = pd.read_excel(excel_path, sheet_name=sheet)
                if not df.empty:
                    sheet_output = f"{base_name}_{sheet}.csv"
                    df.to_csv(sheet_output, index=False)
                    output_files.append(sheet_output)

            return output_files
        else:
            # Convert single sheet
            df = pd.read_excel(excel_path, sheet_name=sheet_name) if sheet_name else pd.read_excel(excel_path)
            df.to_csv(output_path, index=False, encoding='utf-8')
            return output_path
    except Exception as e:
        raise ValueError(f"Error converting Excel to CSV: {str(e)}")


def text_to_html(text_path, **kwargs):
    """Convert plain text to HTML"""
    output_path = kwargs.get('output_path')
    title = kwargs.get('title', 'Converted Document')

    try:
        # Read the text file
        with open(text_path, 'r', encoding='utf-8') as f:
            text_content = f.read()

        # Convert to simple HTML
        paragraphs = text_content.split('\n\n')
        formatted_paragraphs = "".join(f"<p>{paragraph}</p>" for paragraph in paragraphs if paragraph.strip())

        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>{title}</title>
    <meta charset="UTF-8">
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #333; }}
        p {{ margin-bottom: 16px; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    {formatted_paragraphs}
</body>
</html>"""

        # Save HTML file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return output_path
    except Exception as e:
        raise ValueError(f"Error converting text to HTML: {str(e)}")


def image_to_text(image_path, **kwargs):
    """Extract text from image using OCR"""
    output_path = kwargs.get('output_path')

    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)

        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)

        return text
    except Exception as e:
        raise ValueError(f"Error extracting text from image: {str(e)}")


# Register document converters
ConverterFactory.register('docx_to_pdf', docx_to_pdf)
ConverterFactory.register('html_to_pdf', html_to_pdf)
ConverterFactory.register('excel_to_pdf', excel_to_pdf)
ConverterFactory.register('pdf_to_docx', pdf_to_docx)
ConverterFactory.register('create_csv_from_excel', create_csv_from_excel)
ConverterFactory.register('text_to_html', text_to_html)
ConverterFactory.register('image_to_text', image_to_text)